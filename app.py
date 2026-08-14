import streamlit as st
from google import genai
import json
import PyPDF2
import io
import os
import pandas as pd
from datetime import date
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials

# ---------- Page setup ----------
st.set_page_config(page_title="Job Application Assistant Agent", page_icon="💼", layout="wide")

st.markdown("""
<style>
.block-container {
    padding-top: 1rem !important;
    padding-bottom: 1rem !important;
}
.main-header {
    padding: 0.7rem 1.2rem;
    background: linear-gradient(135deg, #2563EB 0%, #1E40AF 100%);
    border-radius: 10px;
    margin-bottom: 1rem;
}
.main-header h1 {
    color: white;
    font-size: 1.2rem;
    margin: 0;
}
.main-header p {
    color: #DBEAFE;
    margin: 0;
    font-size: 0.78rem;
}
.app-item {
    padding: 0.6rem 0.8rem;
    border-radius: 8px;
    margin-bottom: 0.4rem;
    background: white;
    border-left: 3px solid transparent;
}
.app-item:hover {
    background: #EFF6FF;
    border-left: 3px solid #2563EB;
}
.app-company {
    font-weight: 600;
    font-size: 0.85rem;
    color: #1E293B;
}
.app-role {
    font-size: 0.75rem;
    color: #64748B;
}
div[data-testid="stTextArea"] textarea { border-radius: 10px; }
div.stButton > button {
    border-radius: 8px;
    font-weight: 600;
    padding: 0.5rem 1.2rem;
}
div[data-testid="stExpander"] {
    border-radius: 10px;
    border: 1px solid #E2E8F0;
}
</style>
""", unsafe_allow_html=True)

# ---------- Gmail OAuth config ----------
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
REDIRECT_URI = "https://job-hunting-ai.streamlit.app"

def create_oauth_flow():
    client_config = {
        "web": {
            "client_id": st.secrets["GOOGLE_CLIENT_ID"],
            "client_secret": st.secrets["GOOGLE_CLIENT_SECRET"],
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [REDIRECT_URI]
        }
    }
    flow = Flow.from_client_config(
        client_config,
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI
    )
    return flow

# ---------- Handle OAuth redirect ----------
query_params = st.query_params
if 'code' in query_params and 'gmail_credentials' not in st.session_state:
    try:
        flow = create_oauth_flow()
        flow.fetch_token(code=query_params['code'])
        st.session_state['gmail_credentials'] = {
            'token': flow.credentials.token,
            'refresh_token': flow.credentials.refresh_token,
            'token_uri': flow.credentials.token_uri,
            'client_id': flow.credentials.client_id,
            'client_secret': flow.credentials.client_secret,
            'scopes': list(flow.credentials.scopes)
        }
        st.query_params.clear()
    except Exception as e:
        st.error(f"Gmail connection error: {str(e)}")

# ---------- PDF extraction ----------
def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text.strip()

# ---------- Connect to Gemini ----------
api_key = st.secrets["GEMINI_API_KEY"]
client = genai.Client(api_key=api_key)

# ---------- Storage functions ----------
TRACKER_FILE = "applications.json"

def load_applications():
    if os.path.exists(TRACKER_FILE):
        with open(TRACKER_FILE, "r") as f:
            return json.load(f)
    return []

def save_applications(applications):
    with open(TRACKER_FILE, "w") as f:
        json.dump(applications, f, indent=2)

def log_application(company, role, status="Applied", notes=""):
    applications = load_applications()
    applications.append({
        "company": company,
        "role": role,
        "date_applied": str(date.today()),
        "status": status,
        "notes": notes
    })
    save_applications(applications)

def update_application_status(company, role, new_status):
    applications = load_applications()
    for app in applications:
        if app['company'].lower() == company.lower() and \
           app['role'].lower() == role.lower():
            app['status'] = new_status
    save_applications(applications)

# ---------- Cover letter generator ----------
def generate_application_materials(resume_text, job_posting_text):
    prompt = f"""
You are a career assistant agent. You will be given a candidate's resume and a job posting.

Your tasks:
1. Write a tailored, professional cover letter (3-4 paragraphs) that connects the candidate's
   real skills/experience to this specific job posting. Do not invent skills or experience
   that aren't in the resume.
2. After the cover letter, add a section titled "SKILL MATCH" that lists:
   - Matched Skills: skills/experience from the resume that align with the job requirements
   - Gaps: requirements in the job posting that the resume doesn't clearly show

RESUME:
{resume_text}

JOB POSTING:
{job_posting_text}

Format your response with clear headers: "COVER LETTER" and "SKILL MATCH".
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text

# ---------- Export functions ----------
def export_to_pdf(text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20*mm,
        rightMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=20*mm
    )
    styles = getSampleStyleSheet()
    story = []
    for line in text.split('\n'):
        line = line.strip()
        if line == "":
            story.append(Spacer(1, 6))
        elif line.startswith("##") or (line.startswith("**") and line.endswith("**")):
            heading = line.replace("##", "").replace("**", "").strip()
            story.append(Paragraph(f"<b>{heading}</b>", styles['Heading2']))
        elif line.startswith("*") or line.startswith("-"):
            bullet = line.lstrip("*-").strip()
            story.append(Paragraph(f"• {bullet}", styles['Normal']))
        else:
            story.append(Paragraph(line, styles['Normal']))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_docx(text):
    doc = Document()
    doc.add_heading("Cover Letter & Skill Match", 0)
    for line in text.split('\n'):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
        else:
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ---------- Gmail functions ----------
def get_gmail_service():
    creds_data = st.session_state.get('gmail_credentials')
    if not creds_data:
        return None
    creds = Credentials(
        token=creds_data['token'],
        refresh_token=creds_data.get('refresh_token'),
        token_uri=creds_data['token_uri'],
        client_id=creds_data['client_id'],
        client_secret=creds_data['client_secret'],
        scopes=creds_data['scopes']
    )
    return build('gmail', 'v1', credentials=creds)

def scan_emails_for_company(company):
    service = get_gmail_service()
    if not service:
        return []
    try:
        query = f"from:{company.lower()}"
        results = service.users().messages().list(
            userId='me',
            q=query,
            maxResults=5
        ).execute()
        messages = results.get('messages', [])
        emails = []
        for msg in messages:
            msg_data = service.users().messages().get(
                userId='me',
                id=msg['id'],
                format='metadata',
                metadataHeaders=['Subject', 'From', 'Date']
            ).execute()
            headers = msg_data['payload']['headers']
            subject = next(
                (h['value'] for h in headers if h['name'] == 'Subject'), 'No subject'
            )
            snippet = msg_data.get('snippet', '')
            emails.append({
                'company': company,
                'subject': subject,
                'snippet': snippet
            })
        return emails
    except Exception as e:
        st.error(f"Error reading Gmail: {str(e)}")
        return []

def analyze_email_with_gemini(email):
    prompt = f"""
Analyze this email from a company and determine the job application status.

Company: {email['company']}
Email subject: {email['subject']}
Email preview: {email['snippet']}

What is the application status based on this email?
Choose EXACTLY one from:
- Applied
- Interview Scheduled
- Offer
- Rejected

Respond with ONLY the status, nothing else.
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text.strip()

# ---------- Status badge helper ----------
def status_badge(status):
    badges = {
        "Applied": "🔵 Applied",
        "Interview Scheduled": "🟢 Interview",
        "Offer": "🟡 Offer",
        "Rejected": "🔴 Rejected"
    }
    return badges.get(status, status)

# ---------- Sidebar ----------
with st.sidebar:
    st.markdown("## 💼 Job Agent")
    st.button("➕ New Application", use_container_width=True)
    st.markdown("---")
    st.markdown("#### 📋 Applications")
    applications = load_applications()
    if not applications:
        st.caption("No applications yet.")
    else:
        for status_group in ["Applied", "Interview Scheduled", "Offer", "Rejected"]:
            group = [a for a in applications if a['status'] == status_group]
            if group:
                st.markdown(f"**{status_badge(status_group)}**")
                for app in group:
                    st.markdown(f"""
                    <div class="app-item">
                        <div class="app-company">{app['company']}</div>
                        <div class="app-role">{app['role']}</div>
                    </div>
                    """, unsafe_allow_html=True)
    st.markdown("---")
    if 'gmail_credentials' in st.session_state:
        st.success("📧 Gmail Connected")
    else:
        st.caption("📧 Gmail not connected")
    st.write("[GitHub](https://github.com/mujeeb-ullah-khan/job-application-assistant-agent)")
    st.caption("Built by Mujeeb Ullah Khan")

# ---------- Compact header ----------
st.markdown("""
<div class="main-header">
    <h1>💼 Job/Internship Application Assistant Agent</h1>
    <p>AI-powered cover letters, skill matching & application tracking</p>
</div>
""", unsafe_allow_html=True)

# ---------- UI: Tabs ----------
tab1, tab2, tab3, tab4 = st.tabs([
    "✍️ Generate Cover Letter",
    "📋 Application Tracker",
    "📊 Statistics",
    "📧 Gmail Sync"
])

# --- Tab 1: Cover letter generator ---
with tab1:
    st.subheader("Add your resume and a job posting")
    uploaded_pdf = st.file_uploader("📄 Upload Resume (PDF)", type=["pdf"])
    pasted_resume = st.text_area(
        "✍️ Or paste your resume here",
        height=150,
        placeholder="Paste your resume text here..."
    )
    resume_text = ""
    if uploaded_pdf is not None:
        resume_text = extract_text_from_pdf(uploaded_pdf)
        st.success("✅ Resume extracted from PDF!")
        with st.expander("Preview extracted text"):
            st.write(resume_text)
    elif pasted_resume.strip():
        resume_text = pasted_resume
        st.info("Using pasted resume text.")

    job_posting_text = st.text_area(
        "📋 Job Posting",
        height=200,
        placeholder="Paste the job posting here..."
    )

    if st.button("Generate Cover Letter"):
        if resume_text.strip() and job_posting_text.strip():
            with st.spinner("Generating..."):
                result = generate_application_materials(resume_text, job_posting_text)
            st.session_state['generated_result'] = result
        else:
            st.warning("Please fill in both the resume and job posting.")

    if 'generated_result' in st.session_state:
        st.markdown(st.session_state['generated_result'])
        st.divider()
        st.subheader("📥 Export")
        col1, col2 = st.columns(2)
        with col1:
            pdf_bytes = export_to_pdf(st.session_state['generated_result'])
            st.download_button(
                label="⬇️ Download as PDF",
                data=pdf_bytes,
                file_name="cover_letter.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        with col2:
            docx_bytes = export_to_docx(st.session_state['generated_result'])
            st.download_button(
                label="⬇️ Download as DOCX",
                data=docx_bytes,
                file_name="cover_letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    st.divider()
    st.subheader("Log this application")
    col1, col2 = st.columns(2)
    with col1:
        company_input = st.text_input("Company name")
    with col2:
        role_input = st.text_input("Role")
    notes_input = st.text_input("Notes (optional)")
    if st.button("Log Application"):
        if company_input and role_input:
            log_application(company_input, role_input, notes=notes_input)
            st.success(f"Logged: {role_input} at {company_input}")
        else:
            st.warning("Please enter at least company and role.")

# --- Tab 2: Application tracker ---
with tab2:
    st.subheader("Your Applications")
    applications = load_applications()
    if not applications:
        st.info("No applications logged yet.")
    else:
        for i, app in enumerate(applications):
            with st.expander(
                f"{app['role']} at {app['company']} — {status_badge(app['status'])}"
            ):
                st.write(f"**Date applied:** {app['date_applied']}")
                st.write(f"**Notes:** {app['notes'] if app['notes'] else 'None'}")
                new_status = st.selectbox(
                    "Update status",
                    ["Applied", "Interview Scheduled", "Offer", "Rejected"],
                    index=["Applied", "Interview Scheduled", "Offer", "Rejected"].index(
                        app['status']
                    ) if app['status'] in [
                        "Applied", "Interview Scheduled", "Offer", "Rejected"
                    ] else 0,
                    key=f"status_{i}"
                )
                if st.button("Update", key=f"update_{i}"):
                    update_application_status(app['company'], app['role'], new_status)
                    st.success("Updated! Refresh to see changes.")

# --- Tab 3: Statistics ---
with tab3:
    st.subheader("📊 Application Statistics")
    applications = load_applications()
    if not applications:
        st.info("No applications logged yet. Start tracking to see statistics.")
    else:
        total = len(applications)
        applied = len([a for a in applications if a['status'] == "Applied"])
        interview = len([a for a in applications if a['status'] == "Interview Scheduled"])
        offer = len([a for a in applications if a['status'] == "Offer"])
        rejected = len([a for a in applications if a['status'] == "Rejected"])

        col1, col2, col3, col4, col5 = st.columns(5)
        col1.metric("📁 Total", total)
        col2.metric("🔵 Applied", applied)
        col3.metric("🟢 Interview", interview)
        col4.metric("🟡 Offer", offer)
        col5.metric("🔴 Rejected", rejected)

        st.divider()
        df = pd.DataFrame({
            "Status": ["Applied", "Interview", "Offer", "Rejected"],
            "Count": [applied, interview, offer, rejected]
        })
        st.bar_chart(df.set_index("Status"))

        st.divider()
        st.subheader("🕐 Recent Applications")
        recent = sorted(
            applications,
            key=lambda x: x['date_applied'],
            reverse=True
        )[:5]
        for app in recent:
            col1, col2, col3 = st.columns([2, 2, 1])
            col1.write(f"**{app['company']}**")
            col2.write(app['role'])
            col3.write(status_badge(app['status']))

# --- Tab 4: Gmail Sync ---
with tab4:
    st.subheader("📧 Gmail Auto-Sync")

    if 'gmail_credentials' not in st.session_state:
        st.info("""
        Connect your Gmail to automatically detect interview invites,
        rejections, and offers from companies you've applied to.
        The agent will read emails from those companies and suggest
        status updates automatically.
        """)
        st.warning("⚠️ This app is currently in testing mode — only pre-approved Gmail addresses can connect.")

        flow = create_oauth_flow()
        auth_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        st.session_state['oauth_state'] = state
        st.link_button(
            "🔗 Connect Gmail Account",
            auth_url,
            use_container_width=True
        )
    else:
        st.success("✅ Gmail connected successfully!")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔍 Scan All Companies for Updates", use_container_width=True):
                applications = load_applications()
                if not applications:
                    st.warning("No applications logged yet.")
                else:
                    all_emails = []
                    with st.spinner("Scanning your Gmail..."):
                        for app in applications:
                            emails = scan_emails_for_company(app['company'])
                            all_emails.extend(emails)

                    if not all_emails:
                        st.info("No emails found from your tracked companies.")
                    else:
                        st.session_state['scanned_emails'] = all_emails
                        st.success(f"Found {len(all_emails)} emails!")

        with col2:
            if st.button("🔌 Disconnect Gmail", use_container_width=True):
                del st.session_state['gmail_credentials']
                st.rerun()

        if 'scanned_emails' in st.session_state:
            st.divider()
            st.subheader("📬 Emails Found")
            for idx, email in enumerate(st.session_state['scanned_emails']):
                with st.expander(f"{email['company']} — {email['subject'][:60]}"):
                    st.write(f"**Preview:** {email['snippet']}")
                    if st.button(
                        "🤖 Analyze & Detect Status",
                        key=f"analyze_{idx}"
                    ):
                        with st.spinner("Gemini is analyzing the email..."):
                            detected_status = analyze_email_with_gemini(email)
                        st.success(f"Detected status: **{detected_status}**")

                        applications = load_applications()
                        matching = [
                            a for a in applications
                            if a['company'].lower() == email['company'].lower()
                        ]
                        if matching:
                            if st.button(
                                f"✅ Update {email['company']} → {detected_status}",
                                key=f"apply_{idx}"
                            ):
                                update_application_status(
                                    email['company'],
                                    matching[0]['role'],
                                    detected_status
                                )
                                st.success("Status updated!")
                                st.rerun()
